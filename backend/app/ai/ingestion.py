"""Multimodal Ingestion Pipeline.

Handles: PDF text, PDF scan (OCR), tables, images, DOCX.
Includes error handling and fallback strategies.
"""

import os
import io
import logging
from typing import Optional
from PIL import Image

from app.ai.chunking import chunk_legal_text
from app.ai.ocr_processor import ocr_pdf_page
from app.ai.table_extractor import extract_tables_from_pdf
from app.ai.vectorstore import add_chunks_to_store
from app.config import settings

logger = logging.getLogger(__name__)


def process_document(
    file_path: str,
    metadata: Optional[dict] = None,
    collection_name: Optional[str] = None,
) -> dict:
    """Process a document through the full ingestion pipeline.

    Returns:
        {"status": "done"|"error", "text": str, "chunks": int,
         "is_scanned": bool, "tables": int, "error": str|None}
    """
    metadata = metadata or {}
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            return _process_pdf(file_path, metadata, collection_name)
        elif ext == ".docx":
            return _process_docx(file_path, metadata, collection_name)
        else:
            return {"status": "error", "error": f"Unsupported file type: {ext}"}
    except Exception as e:
        logger.error(f"Ingestion failed for {file_path}: {e}")
        return {"status": "error", "error": str(e)}


def _process_pdf(
    file_path: str, metadata: dict, collection_name: Optional[str]
) -> dict:
    """Process PDF: detect text vs scan, extract text, tables, images."""
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    all_text = []
    is_scanned = False
    total_pages = len(doc)

    for page_num in range(total_pages):
        page = doc[page_num]

        # Try text extraction first
        text = page.get_text("text").strip()

        if len(text) < 50:
            # Likely a scanned page — use OCR
            is_scanned = True
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = ocr_pdf_page(img)
            logger.info(f"OCR page {page_num + 1}: {len(text)} chars")

        if text:
            all_text.append(text)

        # Extract images for Vision LLM description
        images = page.get_images(full=True)
        for img_idx, img_info in enumerate(images):
            try:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                if base_image and base_image["image"]:
                    img = Image.open(io.BytesIO(base_image["image"]))
                    # Only process large images (likely diagrams/charts)
                    if img.width > 200 and img.height > 200:
                        desc = _describe_image(img)
                        if desc:
                            all_text.append(f"[Hình ảnh trang {page_num+1}]: {desc}")
            except Exception as e:
                logger.warning(f"Image extraction failed page {page_num+1}: {e}")

    doc.close()

    # Extract tables
    tables = extract_tables_from_pdf(file_path)
    for t in tables:
        all_text.append(f"[Bảng trang {t['page']}]:\n{t['text']}")

    full_text = "\n\n".join(all_text)

    # Chunk and index
    meta = {**metadata, "source_file": os.path.basename(file_path)}
    chunks = chunk_legal_text(full_text, metadata=meta)
    chunk_count = add_chunks_to_store(chunks, collection_name)

    return {
        "status": "done",
        "text": full_text,
        "chunks": chunk_count,
        "is_scanned": is_scanned,
        "tables": len(tables),
        "total_pages": total_pages,
        "error": None,
    }


def _process_docx(
    file_path: str, metadata: dict, collection_name: Optional[str]
) -> dict:
    """Process DOCX files."""
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    # Also extract tables
    table_texts = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_texts.append("\n".join(rows))

    if table_texts:
        full_text += "\n\n" + "\n\n".join(table_texts)

    meta = {**metadata, "source_file": os.path.basename(file_path)}
    chunks = chunk_legal_text(full_text, metadata=meta)
    chunk_count = add_chunks_to_store(chunks, collection_name)

    return {
        "status": "done",
        "text": full_text,
        "chunks": chunk_count,
        "is_scanned": False,
        "tables": len(table_texts),
        "total_pages": 0,
        "error": None,
    }


def _describe_image(image: Image.Image) -> str:
    """Use Gemini Vision to describe an image from a legal document."""
    try:
        import google.generativeai as genai
        from app.core.rate_limiter import gemini_limiter

        genai.configure(api_key=settings.GOOGLE_API_KEY)
        model = genai.GenerativeModel(settings.GEMINI_MODEL)

        gemini_limiter.wait_if_needed()
        from app.ai.prompt_templates import IMAGE_DESCRIPTION_PROMPT

        response = model.generate_content([IMAGE_DESCRIPTION_PROMPT, image])
        return response.text.strip()
    except Exception as e:
        logger.warning(f"Image description failed: {e}")
        return ""
